"""
Daily portfolio report - infographic style (Korean market report style)
Output: agent_Stocks/portfolio_report_YYYY-MM-DD.doc
"""
import io
import json
from datetime import datetime
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import yfinance as yf

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUTPUT_DIR = Path(r"C:\Users\공동환\Desktop\agent_Stocks")
CACHE_DIR  = Path(__file__).parent / "cache"
MY_PORT    = Path(__file__).parent / "my_portfolio.json"
FONT       = "Malgun Gothic"
WEEKDAYS   = "월화수목금토일"

# ── 색상 ────────────────────────────────────────────────
C = {
    "navy":  RGBColor(0x1B, 0x3A, 0x6B),
    "teal":  RGBColor(0x0D, 0x73, 0x77),
    "gold":  RGBColor(0xC8, 0x96, 0x2A),
    "green": RGBColor(0x1A, 0x7F, 0x4B),
    "red":   RGBColor(0xB0, 0x30, 0x30),
    "gray":  RGBColor(0x88, 0x88, 0x88),
    "white": RGBColor(0xFF, 0xFF, 0xFF),
    "black": RGBColor(0x22, 0x22, 0x22),
    "orange": RGBColor(0xE0, 0x70, 0x20),
}
HEX = {
    "navy": "1B3A6B", "teal": "0D7377", "gold": "C8962A",
    "dark": "0F1F3D", "green_d": "0F5C32", "red_d": "7A1A1A",
    "green_l": "D5F5E3", "red_l": "FADBD8", "gold_l": "FEF9E7",
    "blue_l": "D6E4F0", "gray_l": "F2F3F4", "white": "FFFFFF",
    "card1": "EBF5FB", "card2": "FEF9E7",
}

plt.rcParams.update({
    "font.family": "Malgun Gothic",
    "axes.unicode_minus": False,
    "figure.facecolor": "none",
    "axes.facecolor": "none",
})

# ── XML 헬퍼 ────────────────────────────────────────────

def _shade(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.upper())
    tcPr.append(shd)


def _font(run, size=None, bold=False, color=None):
    run.font.name = FONT
    rPr = run._r.get_or_add_rPr()
    rF  = rPr.find(qn("w:rFonts"))
    if rF is None:
        rF = OxmlElement("w:rFonts")
        rPr.insert(0, rF)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rF.set(qn(attr), FONT)
    if size:  run.font.size = Pt(size)
    if bold:  run.bold = True
    if color: run.font.color.rgb = color


def _cwrite(cell, text, size=9, bold=False, color=None,
            align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0):
    para = cell.paragraphs[0]
    para.clear()
    para.alignment = align
    if space_before:
        para.paragraph_format.space_before = Pt(space_before)
    run = para.add_run(text if isinstance(text, str) else str(text))
    _font(run, size=size, bold=bold, color=color)


def _add_run(para, text, size=9, bold=False, color=None):
    run = para.add_run(text if isinstance(text, str) else str(text))
    _font(run, size=size, bold=bold, color=color)
    return run


def _border(table, hex_color="1B3A6B", sz="4"):
    tbl  = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    bdr = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), sz)
        b.set(qn("w:color"), hex_color)
        bdr.append(b)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    tblPr.append(bdr)


def _no_space(para):
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)


def _set_default_font(doc):
    styles_el = doc.styles.element
    doc_def   = styles_el.find(qn("w:docDefaults"))
    if doc_def is None:
        doc_def = OxmlElement("w:docDefaults"); styles_el.insert(0, doc_def)
    rPr_def = doc_def.find(qn("w:rPrDefault"))
    if rPr_def is None:
        rPr_def = OxmlElement("w:rPrDefault"); doc_def.append(rPr_def)
    rPr = rPr_def.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr"); rPr_def.append(rPr)
    for old in rPr.findall(qn("w:rFonts")):
        rPr.remove(old)
    rF = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rF.set(qn(attr), FONT)
    rPr.insert(0, rF)


def _hrow(table, texts, bg=HEX["navy"], size=9):
    row = table.rows[0]
    for i, t in enumerate(texts):
        _shade(row.cells[i], bg)
        _cwrite(row.cells[i], t, size=size, bold=True, color=C["white"])


# ── 차트 생성 ────────────────────────────────────────────

def _line_chart(prices: list, color: str, w=3.2, h=1.0) -> io.BytesIO:
    fig, ax = plt.subplots(figsize=(w, h))
    xs = list(range(len(prices)))
    ax.plot(xs, prices, color=color, linewidth=2.2, solid_capstyle="round")
    ax.fill_between(xs, prices, min(prices) * 0.9995, alpha=0.18, color=color)
    ax.set_xlim(-0.2, len(prices) - 0.8)
    ax.axis("off")
    buf = io.BytesIO()
    plt.savefig(buf, format="png", dpi=130, bbox_inches="tight",
                transparent=True, pad_inches=0.02)
    plt.close(fig)
    buf.seek(0)
    return buf


def _hbar_chart(labels: list, values: list, w=5.5, h=2.2) -> io.BytesIO:
    colors = ["#1A7F4B" if v >= 0 else "#B03030" for v in values]
    fig, ax = plt.subplots(figsize=(w, h))
    bars = ax.barh(labels[::-1], values[::-1], color=colors[::-1],
                   height=0.6, edgecolor="none")
    ax.axvline(0, color="#333333", linewidth=0.8)
    for bar, val in zip(bars, values[::-1]):
        x = bar.get_width()
        ax.text(x + (0.05 * max(abs(v) for v in values) if x >= 0 else
                     -0.05 * max(abs(v) for v in values)),
                bar.get_y() + bar.get_height() / 2,
                f"{val:+.1f}%", va="center",
                ha="left" if x >= 0 else "right", fontsize=9, fontweight="bold",
                color="#1A7F4B" if x >= 0 else "#B03030")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_visible(False)
    ax.tick_params(left=False, labelsize=9)
    ax.set_xlabel("일간 수익률 (%)", fontsize=8)
    plt.tight_layout(pad=0.3)
    buf = io.BytesIO()
    plt.savefig(buf, format="png", dpi=130, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


# ── 데이터 수집 ──────────────────────────────────────────

def _fetch_index(ticker: str) -> dict:
    try:
        fi    = yf.Ticker(ticker).fast_info
        price = float(getattr(fi, "last_price", 0) or 0)
        prev  = float(getattr(fi, "previous_close", price) or price)
        chg   = price - prev
        pct   = chg / prev * 100 if prev else 0
        hist  = yf.Ticker(ticker).history(period="5d")["Close"].tolist()
        return {"price": price, "chg": chg, "pct": pct, "hist": hist or [price]}
    except Exception:
        return {"price": 0, "chg": 0, "pct": 0, "hist": [0]}


def _fetch_portfolio(holdings: list) -> dict:
    result = {}
    for h in holdings:
        t = h["ticker"]
        try:
            fi    = yf.Ticker(t).fast_info
            price = float(getattr(fi, "last_price", 0) or 0)
            prev  = float(getattr(fi, "previous_close", price) or price)
            chg   = price - prev
            pct   = chg / prev * 100 if prev else 0
            cost  = h.get("avg_cost", price)
            total_pnl = (price - cost) / cost * 100 if cost else 0
            result[t] = {"price": price, "chg": chg, "pct": pct,
                         "total_pnl": total_pnl, "shares": h.get("shares", 0)}
        except Exception:
            result[t] = {"price": 0, "chg": 0, "pct": 0, "total_pnl": 0, "shares": 0}
    return result


def _get_rs(ticker: str, rs90: list) -> float:
    rs_map = {s["ticker"]: s.get("rs_rating", 0) for s in rs90}
    if ticker in rs_map:
        return rs_map[ticker]
    # 미등재 종목: 간이 추정
    try:
        hist = yf.Ticker(ticker).history(period="1y")["Close"]
        if len(hist) < 50:
            return 0
        r12 = (hist.iloc[-1] / hist.iloc[0] - 1) * 100
        spx = yf.Ticker("^GSPC").history(period="1y")["Close"]
        spx_r = (spx.iloc[-1] / spx.iloc[0] - 1) * 100
        diff = r12 - spx_r
        return round(min(99, max(1, 50 + diff * 0.5)), 1)
    except Exception:
        return 0


# ── 섹션 빌더 ────────────────────────────────────────────

def _build_header(doc, today, weekday, subtitle):
    t = doc.add_table(rows=2, cols=1)
    _border(t, HEX["dark"], sz="6")
    c0 = t.rows[0].cells[0]
    _shade(c0, HEX["dark"])
    para = c0.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _no_space(para)
    _add_run(para, f"  {today[:4]}년 {today[5:7]}월 {today[8:]}일 ({weekday})  ",
             size=18, bold=True, color=C["white"])
    _add_run(para, " 포트폴리오 일일 리포트 ", size=18, bold=True, color=C["gold"])

    c1 = t.rows[1].cells[0]
    _shade(c1, HEX["navy"])
    _cwrite(c1, subtitle, size=9, color=C["white"])


def _build_index_cards(doc, sp, nq, vix):
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    t = doc.add_table(rows=1, cols=2)
    _border(t, HEX["navy"])

    def _fill_card(cell, name, data, hist_color, extra_label, extra_val):
        _shade(cell, HEX["card1"])
        cell.paragraphs[0].clear()
        # 인덱스 이름
        p1 = cell.add_paragraph()
        _no_space(p1)
        arr = "▲" if data["chg"] >= 0 else "▼"
        arr_col = C["green"] if data["chg"] >= 0 else C["red"]
        _add_run(p1, f" {name}  ", size=10, bold=True, color=C["navy"])
        _add_run(p1, arr, size=10, bold=True, color=arr_col)
        # 지수 값
        p2 = cell.add_paragraph()
        _no_space(p2)
        _add_run(p2, f"  {data['price']:,.2f}", size=14, bold=True, color=C["black"])
        # 변동
        p3 = cell.add_paragraph()
        _no_space(p3)
        chg_col = C["green"] if data["chg"] >= 0 else C["red"]
        _add_run(p3, f"  {data['chg']:+.2f}  ({data['pct']:+.2f}%)",
                 size=9, bold=True, color=chg_col)
        # 차트
        if len(data["hist"]) >= 2:
            chart = _line_chart(data["hist"], "#1A7F4B" if data["chg"] >= 0 else "#B03030")
            p4 = cell.add_paragraph()
            _no_space(p4)
            run = p4.add_run()
            run.add_picture(chart, width=Cm(6.5))
        # 추가 정보
        p5 = cell.add_paragraph()
        _no_space(p5)
        _add_run(p5, f"  {extra_label}: ", size=8, color=C["gray"])
        _add_run(p5, extra_val, size=8, bold=True, color=C["navy"])

    _fill_card(t.rows[0].cells[0], "S&P 500", sp, "#1A7F4B",
               "전일 대비", f"{sp['pct']:+.2f}%")
    vix_level = "낮음" if vix["price"] < 15 else ("보통" if vix["price"] < 25 else "높음")
    _fill_card(t.rows[0].cells[1], "NASDAQ",  nq, "#0D7377",
               f"VIX {vix['price']:.1f}", vix_level)


def _build_portfolio_and_drivers(doc, holdings, port_data, rs90, macro):
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    t = doc.add_table(rows=1, cols=2)
    _border(t, HEX["navy"])

    # ── 왼쪽: 내 포트폴리오 성과
    left = t.rows[0].cells[0]
    _shade(left, HEX["white"])
    left.paragraphs[0].clear()
    _no_space(left.paragraphs[0])
    p_title = left.add_paragraph()
    _no_space(p_title)
    _add_run(p_title, " ■ 내 포트폴리오 당일 성과", size=10, bold=True, color=C["navy"])

    tickers = [h["ticker"] for h in holdings]
    changes = [port_data[t].get("pct", 0) for t in tickers]
    prices  = [port_data[t].get("price", 0) for t in tickers]
    total_pnls = [port_data[t].get("total_pnl", 0) for t in tickers]

    # 미니 표
    inner = left.add_table(rows=len(holdings)+1, cols=4)
    _border(inner, HEX["teal"], sz="2")
    _hrow(inner, ["종목", "현재가", "당일", "수익률"], bg=HEX["teal"], size=8)
    for i, h in enumerate(holdings, 1):
        t2 = h["ticker"]
        d  = port_data[t2]
        row = inner.rows[i]
        chg_col = C["green"] if d["pct"] >= 0 else C["red"]
        pnl_col = C["green"] if d["total_pnl"] >= 0 else C["red"]
        bg = HEX["green_l"] if d["pct"] >= 0 else HEX["red_l"]
        _shade(row.cells[0], HEX["white"])
        _cwrite(row.cells[0], t2, size=8, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
        _cwrite(row.cells[1], f"${d['price']:.1f}", size=8)
        _shade(row.cells[2], bg)
        _cwrite(row.cells[2], f"{d['pct']:+.1f}%", size=8, bold=True, color=chg_col)
        _cwrite(row.cells[3], f"{d['total_pnl']:+.1f}%", size=8, color=pnl_col)

    # 바 차트
    if len(tickers) >= 2:
        chart = _hbar_chart(tickers, changes, w=4.5, h=2.2)
        p_chart = left.add_paragraph()
        _no_space(p_chart)
        p_chart.add_run().add_picture(chart, width=Cm(8))

    # ── 오른쪽: 상승 주도 요인 + RS 강세주
    right = t.rows[0].cells[1]
    _shade(right, HEX["gold_l"])
    right.paragraphs[0].clear()
    _no_space(right.paragraphs[0])

    p_t2 = right.add_paragraph()
    _no_space(p_t2)
    _add_run(p_t2, " ■ 상승 주도 요인", size=10, bold=True, color=C["navy"])

    phase = macro.get("phase", "TRANSITIONAL")
    signals = macro.get("signals", macro.get("macro_signals", {}))
    drivers = {
        "RISK_ON":      ["연준 금리 동결 기대감", "기술주·AI 섹터 강세", "기관 순매수 지속", "달러 약세 유지"],
        "TRANSITIONAL": ["혼조세 장세", "선별적 매수 우세", "실적 발표 주시", "금리 변동 관찰"],
        "RISK_OFF":     ["리스크 오프 심화", "방어주 선호", "기관 매도 우위", "안전자산 수요 증가"],
    }.get(phase, ["시장 분석 진행 중"])

    for drv in drivers:
        p = right.add_paragraph()
        _no_space(p)
        _add_run(p, f"  ● {drv}", size=9, color=C["black"])

    # 시장 강세주 RS Top 5
    p_t3 = right.add_paragraph()
    _no_space(p_t3)
    p_t3.paragraph_format.space_before = Pt(6)
    _add_run(p_t3, " ■ 시장 RS 강세주 Top 5", size=10, bold=True, color=C["navy"])

    rs_top5 = rs90[:5] if rs90 else []
    for i, s in enumerate(rs_top5, 1):
        p = right.add_paragraph()
        _no_space(p)
        _add_run(p, f"  {i}. {s['ticker']:<6}", size=9, bold=True, color=C["navy"])
        _add_run(p, f"  RS {s.get('rs_rating', 0):.1f}", size=9, color=C["teal"])

    # RS 90+ 종목 수
    rs90_cnt = len(rs90)
    p_cnt = right.add_paragraph()
    _no_space(p_cnt)
    p_cnt.paragraph_format.space_before = Pt(4)
    _add_run(p_cnt, f"  상승 종목 (RS≥90): ", size=9, color=C["gray"])
    _add_run(p_cnt, f"{rs90_cnt}개", size=12, bold=True, color=C["navy"])


def _build_canslim(doc, holdings, canslim_scores, rs90):
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    p_t = doc.add_paragraph()
    _no_space(p_t)
    p_t.paragraph_format.space_before = Pt(4)
    _add_run(p_t, "■ CANSLIM 분석  —  내 포트폴리오",
             size=10, bold=True, color=C["navy"])

    cs_map = {s["ticker"]: s.get("canslim_score", s.get("score", 0))
              for s in canslim_scores}
    rs_map = {s["ticker"]: s.get("rs_rating", 0) for s in rs90}

    n = len(holdings)
    t = doc.add_table(rows=n + 1, cols=6)
    _border(t, HEX["navy"])
    _hrow(t, ["종목", "RS 등급", "CANSLIM", "C EPS", "N 고점", "상태"])

    for i, h in enumerate(holdings, 1):
        tk  = h["ticker"]
        rs  = rs_map.get(tk, _get_rs(tk, rs90))
        cs  = cs_map.get(tk, 0)
        row = t.rows[i]

        # 상태 색
        if rs >= 90 and cs >= 50:
            bg, status = HEX["green_l"], "강세 유지"
        elif rs >= 80 and cs >= 40:
            bg, status = HEX["blue_l"],  "보유"
        elif rs >= 60:
            bg, status = HEX["gold_l"],  "주의"
        else:
            bg, status = HEX["red_l"],   "점검 필요"

        _cwrite(row.cells[0], tk, size=9, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
        _cwrite(row.cells[1], f"{rs:.1f}", size=9,
                color=(C["green"] if rs >= 90 else C["gold"] if rs >= 70 else C["red"]))
        _cwrite(row.cells[2], f"{cs}/70", size=9)
        _cwrite(row.cells[3], "—", size=9)    # C criterion placeholder
        _cwrite(row.cells[4], "—", size=9)    # N criterion placeholder
        _shade(row.cells[5], bg)
        _cwrite(row.cells[5], status, size=9, bold=True)


def _build_outlook(doc, holdings, port_data, rs90, phase):
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    p_t = doc.add_paragraph()
    _no_space(p_t)
    p_t.paragraph_format.space_before = Pt(4)
    _add_run(p_t, "■ 향후 전망 및 전략", size=10, bold=True, color=C["navy"])

    rs_map = {s["ticker"]: s.get("rs_rating", 0) for s in rs90}
    cs_map = {s["ticker"]: 0 for s in holdings}

    t = doc.add_table(rows=len(holdings) + 1, cols=5)
    _border(t, HEX["gold"])
    _hrow(t, ["종목", "RS", "추세", "전망 신호", "전략"], bg=HEX["gold"])

    for i, h in enumerate(holdings, 1):
        tk  = h["ticker"]
        rs  = rs_map.get(tk, _get_rs(tk, rs90))
        pct = port_data.get(tk, {}).get("pct", 0)
        row = t.rows[i]

        if rs >= 90 and phase == "RISK_ON":
            sig, sig_hex, sig_col, trend, action = "강력 보유", HEX["green_l"], C["green"], "↑↑", "비중 유지 / 확대 고려"
        elif rs >= 80:
            sig, sig_hex, sig_col, trend, action = "보유",     HEX["blue_l"],  C["navy"],  "↗",  "현 비중 유지"
        elif rs >= 60:
            sig, sig_hex, sig_col, trend, action = "주의",     HEX["gold_l"],  C["gold"],  "→",  "손절 기준 설정, 모니터링"
        else:
            sig, sig_hex, sig_col, trend, action = "축소 검토", HEX["red_l"],   C["red"],   "↘",  "비중 축소 또는 교체"

        _cwrite(row.cells[0], tk, size=9, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
        _cwrite(row.cells[1], f"{rs:.1f}", size=9)
        _cwrite(row.cells[2], trend, size=11, bold=True)
        _shade(row.cells[3], sig_hex)
        _cwrite(row.cells[3], sig, size=9, bold=True, color=sig_col)
        _cwrite(row.cells[4], action, size=8, align=WD_ALIGN_PARAGRAPH.LEFT)


# ── 메인 ─────────────────────────────────────────────────

def build_report(result: dict) -> Document:
    now      = datetime.now()
    today    = now.strftime("%Y-%m-%d")
    weekday  = WEEKDAYS[now.weekday()]

    holdings = (json.loads(MY_PORT.read_text(encoding="utf-8")).get("holdings", [])
                if MY_PORT.exists() else [])

    # 캐시 로드
    rs90 = []
    rs90_cache = CACHE_DIR / "rs90.json"
    if rs90_cache.exists():
        rs90 = json.loads(rs90_cache.read_text(encoding="utf-8")).get("stocks", [])

    canslim_scores = []
    cs_cache = CACHE_DIR / "canslim_top10.json"
    if cs_cache.exists():
        canslim_scores = json.loads(cs_cache.read_text(encoding="utf-8")).get("top10", [])

    macro = {}
    macro_cache = CACHE_DIR / "macro.json"
    if macro_cache.exists():
        macro = json.loads(macro_cache.read_text(encoding="utf-8"))
    else:
        macro = result

    phase = macro.get("phase", result.get("macro_phase", "TRANSITIONAL"))

    # 시장 데이터 수집
    print("  시장 데이터 수집 중...")
    sp  = _fetch_index("^GSPC")
    nq  = _fetch_index("^IXIC")
    vix = _fetch_index("^VIX")

    # 포트폴리오 데이터
    print("  포트폴리오 데이터 수집 중...")
    port_data = _fetch_portfolio(holdings)

    # 자막 생성
    sp_dir  = "급등" if sp["pct"] > 1 else ("상승" if sp["pct"] > 0 else "하락")
    nq_dir  = "급등" if nq["pct"] > 1 else ("상승" if nq["pct"] > 0 else "하락")
    phase_map: dict[str, str] = {"RISK_ON": "강세장", "TRANSITIONAL": "혼조세", "RISK_OFF": "약세장"}
    phase_kr = phase_map.get(str(phase), "—")
    subtitle = (f"S&P500 {sp['pct']:+.2f}% {sp_dir}  ·  "
                f"NASDAQ {nq['pct']:+.2f}% {nq_dir}  ·  "
                f"VIX {vix['price']:.1f}  ·  시장 국면: {phase_kr}")

    doc = Document()
    _set_default_font(doc)
    for sec in doc.sections:
        sec.top_margin    = Cm(1.2)
        sec.bottom_margin = Cm(1.2)
        sec.left_margin   = Cm(1.5)
        sec.right_margin  = Cm(1.5)

    _build_header(doc, today, weekday, subtitle)
    _build_index_cards(doc, sp, nq, vix)
    _build_portfolio_and_drivers(doc, holdings, port_data, rs90, macro)
    _build_canslim(doc, holdings, canslim_scores, rs90)
    _build_outlook(doc, holdings, port_data, rs90, phase)

    # 푸터
    pf = doc.add_paragraph()
    pf.paragraph_format.space_before = Pt(4)
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(pf, "stock-orchestrator 자동 생성  |  본 리포트는 참고용이며 투자 판단의 최종 책임은 투자자 본인에게 있습니다.",
             size=7, color=C["gray"])

    return doc


def save_report(result: dict) -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    today    = datetime.now().strftime("%Y-%m-%d")
    filepath = OUTPUT_DIR / f"portfolio_report_{today}.doc"
    build_report(result).save(str(filepath))
    print(f"  Report saved: {filepath}")
    return filepath


if __name__ == "__main__":
    cache = CACHE_DIR / "latest_result.json"
    result = json.loads(cache.read_text(encoding="utf-8")) if cache.exists() else {}
    save_report(result)

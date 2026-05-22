"""
포트폴리오 분석 리포트
my_portfolio.json -> RS/CANSLIM 분석 -> 시장 비교 -> Word 리포트
Output: agent_Stocks/portfolio_analysis_YYYY-MM-DD.doc
"""
import json
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

sys.path.insert(0, str(Path(__file__).parent))

OUTPUT_DIR  = Path(r"C:\Users\공동환\Desktop\agent_Stocks")
CACHE_DIR   = Path(__file__).parent / "cache"
MY_PORT     = Path(__file__).parent / "my_portfolio.json"
FONT        = "Malgun Gothic"
WEEKDAYS    = "월화수목금토일"

C = {
    "navy":  RGBColor(0x1B, 0x3A, 0x6B),
    "teal":  RGBColor(0x0D, 0x73, 0x77),
    "gold":  RGBColor(0xC8, 0x96, 0x2A),
    "green": RGBColor(0x1A, 0x7F, 0x4B),
    "red":   RGBColor(0xB0, 0x30, 0x30),
    "gray":  RGBColor(0x88, 0x88, 0x88),
    "white": RGBColor(0xFF, 0xFF, 0xFF),
    "black": RGBColor(0x22, 0x22, 0x22),
}
HEX = {
    "navy": "1B3A6B", "teal": "0D7377", "gold": "C8962A",
    "green_light": "D5F5E3", "red_light": "FADBD8",
    "gold_light": "FEF9E7", "blue_light": "D6E4F0",
    "gray_light": "F2F3F4", "white": "FFFFFF",
}

# ── XML helpers ──────────────────────────────────────────

def _shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.upper())
    tcPr.append(shd)


def _font(run, size=None, bold=False, color=None):
    run.font.name = FONT
    rPr = run._r.get_or_add_rPr()
    rF  = rPr.find(qn("w:rFonts"))
    if rF is None:
        rF = OxmlElement("w:rFonts")
        rPr.insert(0, rF)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rF.set(qn(attr), FONT)
    if size:  run.font.size = Pt(size)
    if bold:  run.bold = True
    if color: run.font.color.rgb = color


def _cwrite(cell, text, size=9, bold=False, color=None,
            align=WD_ALIGN_PARAGRAPH.CENTER):
    para = cell.paragraphs[0]
    para.clear()
    para.alignment = align
    run = para.add_run(text if isinstance(text, str) else str(text))
    _font(run, size=size, bold=bold, color=color)


def _border(table, hex_color="1B3A6B"):
    tbl  = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    bdr = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), hex_color)
        bdr.append(b)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    tblPr.append(bdr)


def _set_default_font(doc):
    styles_el = doc.styles.element
    doc_def   = styles_el.find(qn("w:docDefaults"))
    if doc_def is None:
        doc_def = OxmlElement("w:docDefaults")
        styles_el.insert(0, doc_def)
    rPr_def = doc_def.find(qn("w:rPrDefault"))
    if rPr_def is None:
        rPr_def = OxmlElement("w:rPrDefault")
        doc_def.append(rPr_def)
    rPr = rPr_def.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        rPr_def.append(rPr)
    for old in rPr.findall(qn("w:rFonts")):
        rPr.remove(old)
    rF = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rF.set(qn(attr), FONT)
    rPr.insert(0, rF)


def _section_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(f"  {text}")
    _font(run, size=10, bold=True, color=C["navy"])
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    "18")
    left.set(qn("w:color"), HEX["gold"])
    pBdr.append(left)
    pPr.append(pBdr)


def _header_row(table, texts, bg=HEX["navy"]):
    row = table.rows[0]
    for i, text in enumerate(texts):
        _shade(row.cells[i], bg)
        _cwrite(row.cells[i], text, size=9, bold=True, color=C["white"])


# ── 분석 로직 ────────────────────────────────────────────

def load_my_portfolio() -> list[dict]:
    if not MY_PORT.exists():
        return []
    return json.loads(MY_PORT.read_text(encoding="utf-8")).get("holdings", [])


def fetch_current_prices(tickers: list[str]) -> dict[str, float]:
    import yfinance as yf
    prices = {}
    for t in tickers:
        try:
            fi = yf.Ticker(t).fast_info
            prices[t] = float(getattr(fi, "last_price", 0) or 0)
        except Exception:
            prices[t] = 0.0
    return prices


def enrich_holdings(holdings, rs90, canslim_scores, prices) -> list[dict]:
    rs_map = {s["ticker"]: s.get("rs_rating", 0) for s in rs90}
    # 유저 포트폴리오 RS 캐시 병합 (RS90 미만 종목 포함)
    up_file = CACHE_DIR / "user_portfolio_rs.json"
    if up_file.exists():
        for t, r in json.loads(up_file.read_text(encoding="utf-8")).get("ratings", {}).items():
            rs_map.setdefault(t, r)
    cs_map = {s["ticker"]: s.get("canslim_score", s.get("score", 0)) for s in canslim_scores}

    enriched = []
    for h in holdings:
        t       = h["ticker"]
        current = prices.get(t, 0)
        cost    = h.get("avg_cost", 0)
        pnl     = ((current - cost) / cost * 100) if cost else 0
        rs      = rs_map.get(t, 0)
        cs      = cs_map.get(t, 0)

        if rs >= 90 and cs >= 50:
            status, status_hex = "강세 유지", HEX["green_light"]
        elif rs >= 80 and cs >= 40:
            status, status_hex = "보유",     HEX["blue_light"]
        elif rs >= 60:
            status, status_hex = "주의",     HEX["gold_light"]
        else:
            status, status_hex = "점검 필요", HEX["red_light"]

        enriched.append({
            **h,
            "current_price": round(current, 2),
            "pnl_pct":       round(pnl, 1),
            "rs_rating":     round(rs, 1),
            "canslim_score": cs,
            "status":        status,
            "status_hex":    status_hex,
        })
    return enriched


def score_user_canslim(holdings) -> list[dict]:
    """CANSLIM 점수 계산 (캐시 우선, 없으면 직접 계산)."""
    from canslim_analyzer import score_canslim

    canslim_cache = CACHE_DIR / "canslim_top10.json"
    cached = {}
    if canslim_cache.exists():
        for s in json.loads(canslim_cache.read_text(encoding="utf-8")).get("top10", []):
            cached[s["ticker"]] = s

    results = []
    for h in holdings:
        t = h["ticker"]
        if t in cached:
            results.append(cached[t])
        else:
            r = score_canslim(t)
            results.append(r)
    return results


def generate_outlook(h: dict, phase: str) -> dict:
    rs, cs = h.get("rs_rating", 0), h.get("canslim_score", 0)
    pnl    = h.get("pnl_pct", 0)

    if rs >= 90 and cs >= 50 and phase == "RISK_ON":
        return {"signal": "강력 보유", "hex": HEX["green_light"], "color": C["green"],
                "trend": "↑↑", "action": "비중 유지 또는 확대"}
    if rs >= 80 and cs >= 40:
        return {"signal": "보유",     "hex": HEX["blue_light"],  "color": C["navy"],
                "trend": "↗",  "action": "현 비중 유지"}
    if rs >= 60:
        return {"signal": "주의",     "hex": HEX["gold_light"],  "color": C["gold"],
                "trend": "→",  "action": "모니터링 강화, 손절 기준 설정"}
    return {"signal": "축소 검토",   "hex": HEX["red_light"],   "color": C["red"],
            "trend": "↘",  "action": "비중 축소 또는 교체 검토"}


# ── 리포트 빌더 ──────────────────────────────────────────

def build_analysis_report(enriched, rs90, canslim_scores, macro, pipeline_result) -> Document:
    now      = datetime.now()
    today    = now.strftime("%Y-%m-%d")
    weekday  = WEEKDAYS[now.weekday()]
    phase    = macro.get("phase", "TRANSITIONAL")
    signals  = macro.get("signals", macro.get("macro_signals", {}))

    doc = Document()
    _set_default_font(doc)

    for sec in doc.sections:
        sec.top_margin    = Cm(1.5)
        sec.bottom_margin = Cm(1.5)
        sec.left_margin   = Cm(1.8)
        sec.right_margin  = Cm(1.8)

    # ── 배너
    banner_tbl = doc.add_table(rows=2, cols=1)
    _border(banner_tbl, HEX["navy"])
    c0 = banner_tbl.rows[0].cells[0]
    _shade(c0, HEX["navy"])
    _cwrite(c0, "포트폴리오 분석 리포트", size=15, bold=True, color=C["white"])
    c1 = banner_tbl.rows[1].cells[0]
    _shade(c1, HEX["teal"])
    para = c1.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = para.add_run(f"{today} ({weekday})  |  시장: ")
    _font(r1, size=9, color=C["white"])
    phase_labels = {"RISK_ON": "강세", "TRANSITIONAL": "중립", "RISK_OFF": "약세"}
    phase_colors = {"RISK_ON": C["green"], "TRANSITIONAL": C["gold"], "RISK_OFF": C["red"]}
    r2 = para.add_run(phase_labels.get(phase, phase))
    _font(r2, size=9, bold=True, color=phase_colors.get(phase, C["white"]))

    # ── 내 포트폴리오 현황
    _section_title(doc, "내 포트폴리오 현황")
    cols = ["종목", "보유량", "평균단가", "현재가", "수익률", "RS", "CANSLIM", "상태"]
    tbl1 = doc.add_table(rows=len(enriched) + 1, cols=len(cols))
    _border(tbl1, HEX["navy"])
    _header_row(tbl1, cols)

    total_value, total_cost = 0.0, 0.0
    for i, h in enumerate(enriched, 1):
        row = tbl1.rows[i]
        pnl = h["pnl_pct"]
        pnl_color = C["green"] if pnl >= 0 else C["red"]
        vals = [
            h["ticker"],
            str(h.get("shares", "—")),
            f"${h.get('avg_cost', 0):.1f}",
            f"${h['current_price']:.1f}",
            f"{pnl:+.1f}%",
            f"{h['rs_rating']:.1f}",
            f"{h['canslim_score']}/70",
            h["status"],
        ]
        colors = [None, None, None, None, pnl_color, None, None, None]
        for j, (cell, val, col) in enumerate(zip(row.cells, vals, colors)):
            _shade(cell, h["status_hex"] if j == 7 else HEX["white"])
            align = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            _cwrite(cell, val, size=9, color=col, align=align)

        shares = h.get("shares", 0)
        total_value += h["current_price"] * shares
        total_cost  += h.get("avg_cost", 0) * shares

    # 포트폴리오 요약
    total_pnl = ((total_value - total_cost) / total_cost * 100) if total_cost else 0
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    r = p.add_run(
        f"  총 평가액: ${total_value:,.0f}  |  "
        f"총 수익률: {total_pnl:+.1f}%  |  "
        f"보유 종목 수: {len(enriched)}개"
    )
    _font(r, size=9, bold=True, color=C["navy"])

    # ── 시장 강세주 vs 내 포트폴리오 비교
    _section_title(doc, "시장 강세주 vs 내 포트폴리오 비교")
    market_top5 = rs90[:5]
    my_top3     = sorted(enriched, key=lambda x: x.get("rs_rating", 0), reverse=True)[:3]

    tbl2 = doc.add_table(rows=6, cols=4)
    _border(tbl2, HEX["teal"])
    _header_row(tbl2, ["구분", "종목", "RS 등급", "CANSLIM"], bg=HEX["teal"])

    canslim_map = {s["ticker"]: s.get("canslim_score", s.get("score", 0))
                   for s in canslim_scores}

    # 시장 Top 5
    for i, s in enumerate(market_top5[:5], 1):
        r = tbl2.rows[i]
        label = "시장 강세주" if i == 1 else ""
        bg    = HEX["gold_light"] if i <= 3 else HEX["gray_light"]
        _shade(r.cells[0], HEX["blue_light"])
        _cwrite(r.cells[0], label, size=8, color=C["teal"])
        _shade(r.cells[1], bg)
        _cwrite(r.cells[1], s["ticker"], size=9, bold=(i <= 3), align=WD_ALIGN_PARAGRAPH.LEFT)
        _shade(r.cells[2], bg)
        _cwrite(r.cells[2], f"{s.get('rs_rating', 0):.1f}", size=9, bold=(i <= 3))
        _shade(r.cells[3], bg)
        cs_val = canslim_map.get(s["ticker"], s.get("canslim_score", s.get("score", "—")))
        _cwrite(r.cells[3], f"{cs_val}/70" if isinstance(cs_val, int) else "—", size=9)

    # 빈 행 → 내 Top 3
    tbl3 = doc.add_table(rows=4, cols=4)
    _border(tbl3, HEX["navy"])
    _header_row(tbl3, ["구분", "내 종목", "RS 등급", "CANSLIM"])
    for i, h in enumerate(my_top3, 1):
        r = tbl3.rows[i]
        diff = h["rs_rating"] - (market_top5[i - 1]["rs_rating"] if i <= len(market_top5) else 0)
        diff_str = f"({diff:+.1f})" if diff != 0 else ""
        _shade(r.cells[0], HEX["gold_light"])
        _cwrite(r.cells[0], "내 포트폴리오" if i == 1 else "", size=8, color=C["gold"])
        _cwrite(r.cells[1], h["ticker"], size=9, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
        rs_color = C["green"] if h["rs_rating"] >= 90 else (C["gold"] if h["rs_rating"] >= 70 else C["red"])
        _cwrite(r.cells[2], f"{h['rs_rating']:.1f} {diff_str}", size=9, color=rs_color)
        _cwrite(r.cells[3], f"{h['canslim_score']}/70", size=9)

    # ── 향후 전망
    _section_title(doc, "향후 전망 및 전략")
    tbl4 = doc.add_table(rows=len(enriched) + 1, cols=5)
    _border(tbl4, HEX["gold"])
    _header_row(tbl4, ["종목", "RS", "추세", "전망 신호", "전략"], bg=HEX["gold"])
    for i, h in enumerate(enriched, 1):
        out = generate_outlook(h, phase)
        row = tbl4.rows[i]
        _cwrite(row.cells[0], h["ticker"],         size=9, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
        _cwrite(row.cells[1], f"{h['rs_rating']:.1f}", size=9)
        _cwrite(row.cells[2], out["trend"],        size=10, bold=True)
        _shade(row.cells[3], out["hex"])
        _cwrite(row.cells[3], out["signal"],       size=9, bold=True, color=out["color"])
        _cwrite(row.cells[4], out["action"],       size=8, align=WD_ALIGN_PARAGRAPH.LEFT)

    # 거시경제 코멘트
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    vix = signals.get("vix_level", "?")
    y10 = signals.get("yield10y", "?")
    macro_comment = {
        "RISK_ON":      f"현재 강세장 (VIX {vix}, 금리 {y10}%) — 성장주 중심 전략 유지 권장",
        "TRANSITIONAL": f"변동 장세 (VIX {vix}, 금리 {y10}%) — 고RS·고CANSLIM 종목 선별 보유",
        "RISK_OFF":     f"약세장 경고 (VIX {vix}, 금리 {y10}%) — 방어적 포지션 전환 권장",
    }.get(phase, "시장 분석 데이터 없음")
    r = p.add_run(f"  거시 코멘트:  {macro_comment}")
    _font(r, size=9, color=C["navy"])

    # 푸터
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("stock-orchestrator 자동 생성  |  투자 판단의 최종 책임은 투자자 본인에게 있습니다.")
    _font(r2, size=8, color=C["gray"])

    return doc


# ── 진입점 ───────────────────────────────────────────────

def run_portfolio_analysis(pipeline_result: dict) -> Path | None:
    holdings = load_my_portfolio()
    if not holdings:
        print("  my_portfolio.json에 보유 종목을 입력하세요.")
        return None

    # 캐시에서 데이터 로드 (API 재호출 최소화)
    rs90 = pipeline_result.get("final5", [])
    canslim_scores = pipeline_result.get("canslim_top10", [])

    rs90_cache = CACHE_DIR / "rs90.json"
    if rs90_cache.exists():
        rs90 = json.loads(rs90_cache.read_text(encoding="utf-8")).get("stocks", [])

    macro_cache = CACHE_DIR / "macro.json"
    macro = (json.loads(macro_cache.read_text(encoding="utf-8"))
             if macro_cache.exists() else pipeline_result)

    # 사용자 보유 종목 CANSLIM (캐시 우선)
    user_canslim = score_user_canslim(holdings)

    # 현재가 fetch (사용자 종목만 — 최소 API 호출)
    tickers = [h["ticker"] for h in holdings]
    prices  = fetch_current_prices(tickers)

    enriched = enrich_holdings(holdings, rs90, user_canslim, prices)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    today    = datetime.now().strftime("%Y-%m-%d")
    filepath = OUTPUT_DIR / f"portfolio_analysis_{today}.doc"

    doc = build_analysis_report(enriched, rs90, user_canslim, macro, pipeline_result)
    doc.save(str(filepath))
    print(f"  Analysis saved: {filepath}")
    return filepath


if __name__ == "__main__":
    cache = CACHE_DIR / "latest_result.json"
    result = json.loads(cache.read_text(encoding="utf-8")) if cache.exists() else {}
    run_portfolio_analysis(result)
